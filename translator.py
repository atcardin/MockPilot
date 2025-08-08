import yaml
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm

class DocCompliler:
    def __init__(self, doc_path, internal_system, external_system, internal_spec_path, external_spec_path):
        self.doc_path = doc_path
        self.internal_system = internal_system
        self.external_system = external_system
        self.internal_spec_path = internal_spec_path
        self.external_spec_path = external_spec_path
        self.doc = Document(self.doc_path)
        self.internal_spec = self.load_config(self.internal_spec_path)
        self.external_spec = self.load_config(self.external_spec_path)

    def load_config(self, config_path):
        try:
            with open(config_path, "r", encoding="utf‑8") as f:
                return yaml.safe_load(f)
        except FileNotFoundError:
            raise Exception(f"Config file not found at {config_path}!")
        
    def write_cell(self, cell, text, bold):
        p = cell.paragraphs[0]
        p.clear()
        p.style = 'TableBody'  # or 'Normal', 'Title', or any custom style
        r = p.add_run(text)
        r.bold = bold

    def parse_schema(self, schema_name, schema, tables, schema_catalog):
        if schema_name in tables:
            return schema_name  # Already parsed

        tables[schema_name] = {}

        # Resolve $ref
        if "$ref" in schema:
            ref_name = schema["$ref"].split("/")[-1]
            ref_schema = schema_catalog.get(ref_name, {})
            return self.parse_schema(ref_name, ref_schema, tables, schema_catalog)

        schema_type = schema.get('type')

        # Handle arrays
        if schema_type == "array":
            items = schema.get("items", {})
            if "$ref" in items:
                item_schema_name = items["$ref"].split("/")[-1]
                item_schema = schema_catalog.get(item_schema_name, {})
                return f"array[{self.parse_schema(item_schema_name, item_schema, tables, schema_catalog)}]"
            elif "type" in items:
                return f"array[{items['type']}]"
            else:
                return "array[unknown]"

        # Handle primitives
        if schema_type not in ["object", "array"] and "properties" not in schema:
            return schema_type or "unknown"

        # Handle object
        schema_fields = {}

        for field, field_details in schema.get("properties", {}).items():
            base_type = field_details.get("type", "")
            field_format = field_details.get("format", "")
            field_type = f"{base_type}/{field_format}" if field_format else base_type

            # Handle $ref fields
            if "$ref" in field_details:
                ref_name = field_details["$ref"].split("/")[-1]
                ref_schema = schema_catalog.get(ref_name, {})
                field_type = self.parse_schema(ref_name, ref_schema, tables, schema_catalog)

            # Handle array fields
            elif base_type == "array":
                items = field_details.get("items", {})
                item_type = items.get("type", "")
                item_format = items.get("format", "")

                if "$ref" in items:
                    item_ref = items["$ref"].split("/")[-1]
                    item_schema = schema_catalog.get(item_ref, {})
                    item_type = self.parse_schema(item_ref, item_schema, tables, schema_catalog)
                elif item_type == "object" and "properties" in items:
                    nested_name = field
                    self.parse_schema(nested_name, items, tables, schema_catalog)
                    item_type = nested_name
                elif item_format:
                    item_type = f"{item_type}/{item_format}"

                field_type = f"array[{item_type}]"

            # Handle inline object fields
            elif base_type == "object" and "properties" in field_details:
                nested_name = field
                self.parse_schema(nested_name, field_details, tables, schema_catalog)
                field_type = nested_name

            schema_fields[field] = {
                "type": field_type,
                "description": field_details.get("description", ""),
                "required": field in schema.get("required", [])
            }

        # Register the schema
        tables[schema_name] = schema_fields
        
        return schema_name

    def parse_endpoint_body(self, base_name, base_schema, schema_catalog):
        # Each schema shall be a different table
        tables = {}
        
        self.parse_schema(base_name, base_schema, tables, schema_catalog)

        if "$ref" in base_schema:
            ref_name = base_schema["$ref"].split("/")[-1]
            tables[base_name] = tables.pop(ref_name)
        
        return tables

    def add_body_table(self, fields, container_cell=None):
        # Add body tables
        if container_cell:
            table = container_cell.add_table(rows=len(fields) + 1, cols=4)
        else:
            table = self.doc.add_table(rows=len(fields) + 1, cols=4)

        table.autofit = False
        table.style = "MessageBody"

        # column_widths = [Cm(3.5), Cm(7.7), Cm(3.8), Cm(2.3)]
        column_widths = [Cm(4.0), Cm(7.2), Cm(3.8), Cm(2.3)]
        headings = ["Property", "Description", "Type/Format", "Required"]

        # Header row
        for i, cell in enumerate(table.rows[0].cells):
            self.write_cell(cell, headings[i], True)
            cell.width = column_widths[i]

        # Data rows
        values = [
            [name, properties["description"], properties["type"], str(properties["required"])]
            for name, properties in fields.items()
        ]

        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            for j, cell in enumerate(row.cells):
                self.write_cell(cell, values[i - 1][j], False)
                cell.width = column_widths[j]


    def add_endpoint(self, sender, receiver, path, method, details, schema_catalog):
        # Endpoint header
        self.doc.add_heading(details.get("operationId", f"{method} {path}"), level=3)
        # Add header table
        table = self.doc.add_table(rows=6, cols=2)
        table.autofit = False
        table.style = "MessageHeader"
        column_widths = [Cm(4.0), Cm(13.3)]
        # column_widths = [Cm(3.5), Cm(13.8)]
        headings = ["Sender","Receiver","HTTP Verb","Endpoint","Trigger","Notes"]
        values = [
            sender,
            receiver,
            method,
            path,
            # details.get("x-trigger", ""),
            details.get("summary", ""),
            details.get("description", "")
        ]
        
        for i, row in enumerate(table.rows):
            self.write_cell(row.cells[0], headings[i], False)
            row.cells[0].width = column_widths[0]
            self.write_cell(row.cells[1], values[i], False)
            row.cells[1].width = column_widths[1]
        
        # If GET method no body is allowed
        if method == "GET":
            # Add sub
            p = self.doc.add_paragraph()
            p.style = "Subheader"
            r = p.add_run("No body required")
            r.bold = True
            return
        # Body needed
        base_schema = details.get("requestBody", {}).get("content",{}).get("application/json",{}).get("schema",{})
        tables = self.parse_endpoint_body("JSON Message Body", base_schema, schema_catalog)

        for parent, fields in tables.items():
            # Create a 1×1 table to keep subheader + table together
            wrapper_table = self.doc.add_table(rows=1, cols=1)
            wrapper_table.autofit = False
            wrapper_table.style = "Wrapper"
            cell = wrapper_table.cell(0, 0)
            cell.width = Cm(18.2)
            p = cell.paragraphs[0]
            p.style = "Subheader"
            r = p.add_run(parent)
            r.bold = True
            self.add_body_table(fields, container_cell = cell)

        # Add response
        base_schema = details.get("responses", {}).get("200",{}).get("content",{}).get("application/json",{}).get("schema",{})

        tables = self.parse_endpoint_body("Response", base_schema, schema_catalog)

        for parent, fields in tables.items():
            wrapper_table = self.doc.add_table(rows=1, cols=1)
            wrapper_table.autofit = False
            wrapper_table.style = "Wrapper"
            cell = wrapper_table.cell(0, 0)
            p = cell.paragraphs[0]
            p.style = "Subheader"
            r = p.add_run(parent)
            r.bold = True
            self.add_body_table(fields, container_cell = cell)


    def add_section(self, sender, receiver, api_specification):
        # Add a page break
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        # Add heading for the section
        self.doc.add_heading(f"Messages from {sender} to {receiver}", level=2)
        # Extract schemas from spec
        schema_catalog = api_specification.get("components",{}).get("schemas",{})
        # Add endpoints
        for path, methods in api_specification.get("paths", {}).items():
            for method, details in methods.items():
                self.add_endpoint(sender, receiver, path, method, details, schema_catalog)
                self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def compile(self):
        #TODO: replace tags
        # From external system to internal section
        self.add_section(self.external_system, self.internal_system, self.internal_spec)
        # From internal system to external section
        self.add_section(self.internal_system, self.external_system, self.external_spec)
        # Save the compiled file
        self.doc.save(f"{self.external_system}-{self.internal_system}-Interface.docx")

if __name__ == "__main__":

    dc = DocCompliler(
        "./word_templates/Template.docx", 
        "Q3-DYMS",
        "EBS",
        "./sender_config.yaml",
        "./receiver_config.yaml")
    
    dc.compile()


